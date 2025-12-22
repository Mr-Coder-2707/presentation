"""
Dataset Loader Module
Supports loading documents from various dataset formats:
- CSV files
- JSON files
- Standard IR datasets (TREC, Cranfield, etc.)
- Uploaded files (TXT, PDF, DOCX)
"""

import os
import json
import csv
from typing import List, Dict, Optional, Tuple, Any
from pathlib import Path
from dataclasses import dataclass, field


@dataclass
class Document:
    """Represents a document with metadata"""
    doc_id: int
    title: str
    content: str
    file_path: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)

    def __repr__(self):
        title_preview = self.title[:30] + '...' if len(self.title) > 30 else self.title
        return f"Document(id={self.doc_id}, title='{title_preview}')"


@dataclass
class RelevanceJudgment:
    """Represents a relevance judgment for evaluation"""
    query_id: str
    doc_id: int
    relevance: int  # 0 = not relevant, 1+ = relevant (higher = more relevant)


@dataclass
class IRDataset:
    """Represents a complete IR dataset with documents, queries, and relevance judgments"""
    name: str
    documents: List[Document]
    queries: Dict[str, str] = field(default_factory=dict)  # query_id -> query text
    relevance_judgments: List[RelevanceJudgment] = field(default_factory=list)

    def get_relevant_docs(self, query_id: str) -> set:
        """Get relevant document IDs for a query"""
        return {rj.doc_id for rj in self.relevance_judgments
                if rj.query_id == query_id and rj.relevance > 0}


class DatasetLoader:
    """
    Comprehensive dataset loader supporting multiple formats
    """

    SUPPORTED_FILE_FORMATS = ['.txt', '.pdf', '.docx', '.doc']
    SUPPORTED_DATASET_FORMATS = ['.csv', '.json', '.jsonl']

    def __init__(self):
        self.documents: List[Document] = []
        self.doc_counter = 0
        self.current_dataset: Optional[IRDataset] = None

    # ==================== File Loading ====================

    def load_text_file(self, file_path: str) -> str:
        """Load content from a text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
            except:
                return ""

    def load_pdf_file(self, file_path: str) -> str:
        """Load content from a PDF file"""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            return text
        except ImportError:
            print("PyPDF2 not installed. Install with: pip install PyPDF2")
            return ""
        except Exception as e:
            print(f"Error reading PDF {file_path}: {e}")
            return ""

    def load_docx_file(self, file_path: str) -> str:
        """Load content from a DOCX file"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except ImportError:
            print("python-docx not installed. Install with: pip install python-docx")
            return ""
        except Exception as e:
            print(f"Error reading DOCX {file_path}: {e}")
            return ""

    def load_single_file(self, file_path: str) -> Optional[Document]:
        """Load a single file and return a Document object"""
        path = Path(file_path)

        if not path.exists():
            print(f"File not found: {file_path}")
            return None

        ext = path.suffix.lower()

        if ext == '.txt':
            content = self.load_text_file(file_path)
        elif ext == '.pdf':
            content = self.load_pdf_file(file_path)
        elif ext in ['.docx', '.doc']:
            content = self.load_docx_file(file_path)
        else:
            print(f"Unsupported format: {ext}")
            return None

        if content and content.strip():
            self.doc_counter += 1
            doc = Document(
                doc_id=self.doc_counter,
                title=path.stem,
                content=content,
                file_path=str(path.absolute()),
                metadata={'format': ext, 'size': len(content)}
            )
            self.documents.append(doc)
            return doc
        return None

    def load_directory(self, directory_path: str, recursive: bool = True) -> List[Document]:
        """Load all supported files from a directory"""
        loaded = []
        path = Path(directory_path)

        if not path.exists():
            print(f"Directory not found: {directory_path}")
            return loaded

        pattern = '**/*' if recursive else '*'

        for file_path in path.glob(pattern):
            if file_path.is_file():
                ext = file_path.suffix.lower()
                if ext in self.SUPPORTED_FILE_FORMATS:
                    doc = self.load_single_file(str(file_path))
                    if doc:
                        loaded.append(doc)
                        print(f"  Loaded: {file_path.name}")

        return loaded

    # ==================== Dataset Loading ====================

    def load_csv_dataset(self, file_path: str,
                         content_column: str = 'content',
                         title_column: str = 'title',
                         id_column: str = None,
                         delimiter: str = ',') -> List[Document]:
        """
        Load documents from a CSV file

        Args:
            file_path: Path to CSV file
            content_column: Name of column containing document text
            title_column: Name of column containing document title
            id_column: Name of column containing document ID (optional)
            delimiter: CSV delimiter character
        """
        loaded = []

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f, delimiter=delimiter)

                for row in reader:
                    content = row.get(content_column, '')
                    title = row.get(title_column, f'Document {self.doc_counter + 1}')

                    if not content:
                        continue

                    self.doc_counter += 1
                    doc_id = int(row[id_column]) if id_column and id_column in row else self.doc_counter

                    # Store all columns as metadata
                    metadata = {k: v for k, v in row.items()
                               if k not in [content_column, title_column, id_column]}

                    doc = Document(
                        doc_id=doc_id,
                        title=title,
                        content=content,
                        file_path=file_path,
                        metadata=metadata
                    )
                    self.documents.append(doc)
                    loaded.append(doc)

            print(f"Loaded {len(loaded)} documents from CSV: {file_path}")

        except Exception as e:
            print(f"Error loading CSV {file_path}: {e}")

        return loaded

    def load_json_dataset(self, file_path: str,
                          content_field: str = 'content',
                          title_field: str = 'title',
                          id_field: str = 'id',
                          documents_key: str = None) -> List[Document]:
        """
        Load documents from a JSON file

        Args:
            file_path: Path to JSON file
            content_field: Field name for document text
            title_field: Field name for document title
            id_field: Field name for document ID
            documents_key: Key containing document array (if nested)
        """
        loaded = []

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            # Handle nested structure
            if documents_key and documents_key in data:
                docs_data = data[documents_key]
            elif isinstance(data, list):
                docs_data = data
            else:
                docs_data = [data]

            for item in docs_data:
                content = item.get(content_field, '')
                title = item.get(title_field, f'Document {self.doc_counter + 1}')

                if not content:
                    continue

                self.doc_counter += 1
                doc_id = item.get(id_field, self.doc_counter)
                if isinstance(doc_id, str):
                    doc_id = self.doc_counter

                # Store remaining fields as metadata
                metadata = {k: v for k, v in item.items()
                           if k not in [content_field, title_field, id_field]}

                doc = Document(
                    doc_id=doc_id,
                    title=title,
                    content=content,
                    file_path=file_path,
                    metadata=metadata
                )
                self.documents.append(doc)
                loaded.append(doc)

            print(f"Loaded {len(loaded)} documents from JSON: {file_path}")

        except Exception as e:
            print(f"Error loading JSON {file_path}: {e}")

        return loaded

    def load_jsonl_dataset(self, file_path: str,
                           content_field: str = 'content',
                           title_field: str = 'title',
                           id_field: str = 'id') -> List[Document]:
        """Load documents from a JSON Lines file (one JSON object per line)"""
        loaded = []

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if not line:
                        continue

                    item = json.loads(line)
                    content = item.get(content_field, '')
                    title = item.get(title_field, f'Document {self.doc_counter + 1}')

                    if not content:
                        continue

                    self.doc_counter += 1
                    doc_id = item.get(id_field, self.doc_counter)

                    doc = Document(
                        doc_id=doc_id if isinstance(doc_id, int) else self.doc_counter,
                        title=title,
                        content=content,
                        file_path=file_path,
                        metadata={k: v for k, v in item.items()
                                 if k not in [content_field, title_field, id_field]}
                    )
                    self.documents.append(doc)
                    loaded.append(doc)

            print(f"Loaded {len(loaded)} documents from JSONL: {file_path}")

        except Exception as e:
            print(f"Error loading JSONL {file_path}: {e}")

        return loaded

    # ==================== Standard IR Datasets ====================

    def load_trec_dataset(self, docs_path: str,
                          queries_path: str = None,
                          qrels_path: str = None) -> IRDataset:
        """
        Load TREC-style dataset

        Format:
        - Documents: <DOC><DOCNO>id</DOCNO><TEXT>content</TEXT></DOC>
        - Queries: query_id query_text
        - Qrels: query_id iter doc_id relevance
        """
        import re

        documents = []
        queries = {}
        relevance_judgments = []

        # Load documents
        if os.path.exists(docs_path):
            with open(docs_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            # Parse TREC format
            doc_pattern = r'<DOC>(.*?)</DOC>'
            docno_pattern = r'<DOCNO>(.*?)</DOCNO>'
            text_pattern = r'<TEXT>(.*?)</TEXT>'

            for match in re.finditer(doc_pattern, content, re.DOTALL):
                doc_content = match.group(1)

                docno_match = re.search(docno_pattern, doc_content)
                text_match = re.search(text_pattern, doc_content, re.DOTALL)

                if docno_match and text_match:
                    self.doc_counter += 1
                    doc = Document(
                        doc_id=self.doc_counter,
                        title=docno_match.group(1).strip(),
                        content=text_match.group(1).strip(),
                        metadata={'original_id': docno_match.group(1).strip()}
                    )
                    documents.append(doc)
                    self.documents.append(doc)

        # Load queries
        if queries_path and os.path.exists(queries_path):
            with open(queries_path, 'r', encoding='utf-8') as f:
                for line in f:
                    parts = line.strip().split(None, 1)
                    if len(parts) == 2:
                        queries[parts[0]] = parts[1]

        # Load relevance judgments
        if qrels_path and os.path.exists(qrels_path):
            with open(qrels_path, 'r', encoding='utf-8') as f:
                for line in f:
                    parts = line.strip().split()
                    if len(parts) >= 4:
                        relevance_judgments.append(RelevanceJudgment(
                            query_id=parts[0],
                            doc_id=int(parts[2]),
                            relevance=int(parts[3])
                        ))

        dataset = IRDataset(
            name="TREC Dataset",
            documents=documents,
            queries=queries,
            relevance_judgments=relevance_judgments
        )

        self.current_dataset = dataset
        print(f"Loaded TREC dataset: {len(documents)} docs, {len(queries)} queries")
        return dataset

    def create_sample_dataset(self) -> IRDataset:
        """Create a sample dataset for testing"""
        sample_docs = [
            ("Introduction to Machine Learning",
             "Machine learning is a subset of artificial intelligence that enables systems to learn and improve from experience. "
             "It focuses on developing computer programs that can access data and use it to learn for themselves. "
             "The process begins with observations or data, such as examples, direct experience, or instruction. "
             "Machine learning algorithms build mathematical models based on sample data to make predictions."),

            ("Deep Learning Fundamentals",
             "Deep learning is part of a broader family of machine learning methods based on artificial neural networks. "
             "Learning can be supervised, semi-supervised or unsupervised. Deep learning architectures such as deep neural networks, "
             "recurrent neural networks, and convolutional neural networks have been applied to fields including computer vision. "
             "Neural networks with many layers can learn complex patterns in large amounts of data."),

            ("Natural Language Processing",
             "Natural language processing (NLP) is a subfield of linguistics, computer science, and artificial intelligence "
             "concerned with the interactions between computers and human language. NLP involves programming computers to process "
             "and analyze large amounts of natural language data including text and speech recognition. "
             "Applications include machine translation, sentiment analysis, and chatbots."),

            ("Information Retrieval Systems",
             "Information retrieval is the process of obtaining relevant information from a collection of resources. "
             "Search engines are the most visible IR applications. The goal is to retrieve documents that are relevant to a user query. "
             "Key metrics include precision, recall, and F1-score for evaluating retrieval effectiveness. "
             "Modern search engines use inverted indexes and ranking algorithms like BM25."),

            ("Database Management Systems",
             "A database management system (DBMS) is software for creating and managing databases. "
             "It provides users and programmers with a systematic way to create, retrieve, update and manage data. "
             "Popular DBMS include MySQL, PostgreSQL, Oracle, and MongoDB for different use cases. "
             "SQL is the standard language for relational database management systems."),

            ("Web Development Technologies",
             "Web development encompasses many types of web content creation including web design, front-end development, "
             "back-end development, and full-stack development. Technologies include HTML, CSS, JavaScript, Python, "
             "and various frameworks like React, Django, and Node.js for building modern web applications. "
             "RESTful APIs enable communication between web services."),

            ("Data Science and Analytics",
             "Data science combines domain expertise, programming skills, and knowledge of mathematics and statistics "
             "to extract meaningful insights from data. Data scientists use machine learning algorithms, data mining techniques, "
             "and big data technologies to analyze and interpret complex data sets for business intelligence. "
             "Python and R are popular programming languages for data science."),

            ("Cloud Computing Fundamentals",
             "Cloud computing is the on-demand availability of computer system resources, especially data storage and computing power. "
             "Major providers include Amazon Web Services, Microsoft Azure, and Google Cloud Platform. "
             "Cloud services include Infrastructure as a Service (IaaS), Platform as a Service (PaaS), and Software as a Service (SaaS). "
             "Containerization with Docker and Kubernetes enables scalable deployments."),

            ("Cybersecurity Principles",
             "Cybersecurity is the practice of protecting systems, networks, and programs from digital attacks. "
             "These attacks are usually aimed at accessing, changing, or destroying sensitive information. "
             "Key concepts include encryption, firewalls, intrusion detection systems, and security protocols. "
             "Authentication and authorization are fundamental to access control."),

            ("Python Programming Language",
             "Python is a high-level, interpreted programming language known for its clear syntax and readability. "
             "It supports multiple programming paradigms including procedural, object-oriented, and functional programming. "
             "Python is widely used in web development, data science, artificial intelligence, and automation scripts. "
             "Popular libraries include NumPy, Pandas, TensorFlow, and Django.")
        ]

        documents = []
        for title, content in sample_docs:
            self.doc_counter += 1
            doc = Document(
                doc_id=self.doc_counter,
                title=title,
                content=content
            )
            documents.append(doc)
            self.documents.append(doc)

        # Sample queries with relevance judgments
        queries = {
            'q1': 'machine learning algorithms',
            'q2': 'deep learning neural networks',
            'q3': 'natural language processing NLP',
            'q4': 'information retrieval search',
            'q5': 'database SQL',
            'q6': 'web development',
            'q7': 'data science analytics',
            'q8': 'cloud computing',
            'q9': 'cybersecurity',
            'q10': 'python programming'
        }

        # Relevance judgments (query_id, doc_id, relevance)
        judgments_data = [
            ('q1', 1, 2), ('q1', 2, 1), ('q1', 7, 1),
            ('q2', 2, 2), ('q2', 1, 1),
            ('q3', 3, 2),
            ('q4', 4, 2),
            ('q5', 5, 2),
            ('q6', 6, 2),
            ('q7', 7, 2), ('q7', 1, 1), ('q7', 2, 1),
            ('q8', 8, 2),
            ('q9', 9, 2),
            ('q10', 10, 2), ('q10', 6, 1), ('q10', 7, 1),
        ]

        relevance_judgments = [
            RelevanceJudgment(query_id=qid, doc_id=did, relevance=rel)
            for qid, did, rel in judgments_data
        ]

        dataset = IRDataset(
            name="Sample IR Dataset",
            documents=documents,
            queries=queries,
            relevance_judgments=relevance_judgments
        )

        self.current_dataset = dataset
        print(f"Created sample dataset: {len(documents)} documents, {len(queries)} queries")
        return dataset

    # ==================== Utility Methods ====================

    def add_document(self, title: str, content: str, metadata: Dict = None) -> Document:
        """Add a document manually"""
        self.doc_counter += 1
        doc = Document(
            doc_id=self.doc_counter,
            title=title,
            content=content,
            metadata=metadata or {}
        )
        self.documents.append(doc)
        return doc

    def get_all_documents(self) -> List[Document]:
        """Get all loaded documents"""
        return self.documents

    def get_document(self, doc_id: int) -> Optional[Document]:
        """Get document by ID"""
        for doc in self.documents:
            if doc.doc_id == doc_id:
                return doc
        return None

    def clear(self):
        """Clear all loaded documents"""
        self.documents = []
        self.doc_counter = 0
        self.current_dataset = None

    def get_statistics(self) -> Dict:
        """Get dataset statistics"""
        if not self.documents:
            return {'total_documents': 0}

        total_content_length = sum(len(doc.content) for doc in self.documents)

        return {
            'total_documents': len(self.documents),
            'total_content_length': total_content_length,
            'avg_document_length': total_content_length / len(self.documents),
            'has_dataset': self.current_dataset is not None,
            'dataset_name': self.current_dataset.name if self.current_dataset else None,
            'num_queries': len(self.current_dataset.queries) if self.current_dataset else 0
        }

