"""
Document Loader Module
Handles loading and preprocessing of various document formats (txt, pdf, docx)
"""

from typing import List, Optional
from pathlib import Path


class Document:
    """Represents a document with its content and metadata"""

    def __init__(self, doc_id: int, title: str, content: str, file_path: str = None):
        self.doc_id = doc_id
        self.title = title
        self.content = content
        self.file_path = file_path
        self.preprocessed_content = ""

    def __repr__(self):
        return f"Document(id={self.doc_id}, title='{self.title[:30]}...')"


class DocumentLoader:
    """Loads documents from various file formats"""

    SUPPORTED_FORMATS = ['.txt', '.pdf', '.docx', '.doc']

    def __init__(self):
        self.documents: List[Document] = []
        self.doc_counter = 0

    def load_text_file(self, file_path: str) -> str:
        """Load content from a text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()

    def load_pdf_file(self, file_path: str) -> str:
        """Load content from a PDF file"""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except ImportError:
            print("PyPDF2 not installed. Install with: pip install PyPDF2")
            return ""
        except Exception as e:
            print(f"Error reading PDF {file_path}: {e}")
            return ""

    def load_docx_file(self, file_path: str) -> str:
        """Load content from a DOCX file"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except ImportError:
            print("python-docx not installed. Install with: pip install python-docx")
            return ""
        except Exception as e:
            print(f"Error reading DOCX {file_path}: {e}")
            return ""

    def load_file(self, file_path: str) -> Optional[Document]:
        """Load a single file and return a Document object"""
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == '.txt':
            content = self.load_text_file(file_path)
        elif ext == '.pdf':
            content = self.load_pdf_file(file_path)
        elif ext in ['.docx', '.doc']:
            content = self.load_docx_file(file_path)
        else:
            print(f"Unsupported format: {ext}")
            return None

        if content:
            self.doc_counter += 1
            doc = Document(
                doc_id=self.doc_counter,
                title=path.name,
                content=content,
                file_path=str(path.absolute())
            )
            self.documents.append(doc)
            return doc
        return None

    def load_directory(self, directory_path: str, recursive: bool = True) -> List[Document]:
        """Load all supported documents from a directory"""
        loaded_docs = []
        path = Path(directory_path)

        if not path.exists():
            print(f"Directory not found: {directory_path}")
            return loaded_docs

        pattern = '**/*' if recursive else '*'

        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_FORMATS:
                doc = self.load_file(str(file_path))
                if doc:
                    loaded_docs.append(doc)
                    print(f"Loaded: {file_path.name}")

        return loaded_docs

    def add_text_document(self, title: str, content: str) -> Document:
        """Add a document directly from text"""
        self.doc_counter += 1
        doc = Document(
            doc_id=self.doc_counter,
            title=title,
            content=content
        )
        self.documents.append(doc)
        return doc

    def add_sample_documents(self) -> List[Document]:
        """Add sample documents for testing"""
        sample_docs = [
            ("Introduction to Machine Learning",
             "Machine learning is a subset of artificial intelligence that enables systems to learn and improve from experience. "
             "It focuses on developing computer programs that can access data and use it to learn for themselves. "
             "The process begins with observations or data, such as examples, direct experience, or instruction."),

            ("Deep Learning Fundamentals",
             "Deep learning is part of a broader family of machine learning methods based on artificial neural networks. "
             "Learning can be supervised, semi-supervised or unsupervised. Deep learning architectures such as deep neural networks, "
             "recurrent neural networks, and convolutional neural networks have been applied to fields including computer vision."),

            ("Natural Language Processing",
             "Natural language processing (NLP) is a subfield of linguistics, computer science, and artificial intelligence "
             "concerned with the interactions between computers and human language. NLP involves programming computers to process "
             "and analyze large amounts of natural language data including text and speech recognition."),

            ("Information Retrieval Systems",
             "Information retrieval is the process of obtaining relevant information from a collection of resources. "
             "Search engines are the most visible IR applications. The goal is to retrieve documents that are relevant to a user query. "
             "Key metrics include precision, recall, and F1-score for evaluating retrieval effectiveness."),

            ("Database Management Systems",
             "A database management system (DBMS) is software for creating and managing databases. "
             "It provides users and programmers with a systematic way to create, retrieve, update and manage data. "
             "Popular DBMS include MySQL, PostgreSQL, Oracle, and MongoDB for different use cases."),

            ("Web Development Technologies",
             "Web development encompasses many types of web content creation including web design, front-end development, "
             "back-end development, and full-stack development. Technologies include HTML, CSS, JavaScript, Python, "
             "and various frameworks like React, Django, and Node.js for building modern web applications."),

            ("Data Science and Analytics",
             "Data science combines domain expertise, programming skills, and knowledge of mathematics and statistics "
             "to extract meaningful insights from data. Data scientists use machine learning algorithms, data mining techniques, "
             "and big data technologies to analyze and interpret complex data sets for business intelligence."),

            ("Cloud Computing Fundamentals",
             "Cloud computing is the on-demand availability of computer system resources, especially data storage and computing power. "
             "Major providers include Amazon Web Services, Microsoft Azure, and Google Cloud Platform. "
             "Cloud services include Infrastructure as a Service (IaaS), Platform as a Service (PaaS), and Software as a Service (SaaS)."),

            ("Cybersecurity Principles",
             "Cybersecurity is the practice of protecting systems, networks, and programs from digital attacks. "
             "These attacks are usually aimed at accessing, changing, or destroying sensitive information. "
             "Key concepts include encryption, firewalls, intrusion detection systems, and security protocols."),

            ("Python Programming Language",
             "Python is a high-level, interpreted programming language known for its clear syntax and readability. "
             "It supports multiple programming paradigms including procedural, object-oriented, and functional programming. "
             "Python is widely used in web development, data science, artificial intelligence, and automation scripts.")
        ]

        docs = []
        for title, content in sample_docs:
            doc = self.add_text_document(title, content)
            docs.append(doc)

        return docs

    def get_all_documents(self) -> List[Document]:
        """Return all loaded documents"""
        return self.documents

    def clear_documents(self):
        """Clear all documents"""
        self.documents = []
        self.doc_counter = 0
